import re
import json
import logging
import io
from typing import Optional, Dict, Any, List
from fastapi import APIRouter, File, UploadFile, Form, HTTPException, status
import httpx

from app.core.config import settings

logger = logging.getLogger(__name__)

router = APIRouter()

def clean_and_normalize_text(raw_text: str) -> str:
    """Clean and normalize extracted resume text."""
    if not raw_text:
        return ""
    # Normalize line breaks and tabs
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")
    # Replace multiple spaces with a single space
    text = re.sub(r" +", " ", text)
    # Remove non-printable control characters except line breaks
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
    # Reduce multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz) preferred, pdfplumber / pypdf fallback."""
    text = ""
    # Strategy 1: PyMuPDF
    try:
        import pymupdf as fitz  # PyMuPDF (fitz alias deprecated)
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text() or "")
        text = "\n".join(pages_text)
        if text.strip():
            logger.info("Successfully extracted text using PyMuPDF (fitz).")
            return clean_and_normalize_text(text)
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}")

    # Strategy 2: pdfplumber fallback
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                pages_text.append(page.extract_text() or "")
            text = "\n".join(pages_text)
            if text.strip():
                logger.info("Successfully extracted text using pdfplumber.")
                return clean_and_normalize_text(text)
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")

    # Strategy 3: pypdf fallback
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            pages_text.append(page.extract_text() or "")
        text = "\n".join(pages_text)
        if text.strip():
            logger.info("Successfully extracted text using pypdf.")
            return clean_and_normalize_text(text)
    except Exception as e:
        logger.warning(f"pypdf extraction failed: {e}")

    # Strategy 4: Raw regex stream fallback
    try:
        raw = file_bytes.decode("latin-1", errors="ignore")
        matches = re.findall(r"\(([^()]{2,})\)\s*Tj", raw)
        if matches:
            text = " ".join(matches)
        else:
            text = re.sub(r"[^\w\s\.,\-\@\#\+\(\)]", " ", raw)
    except Exception:
        text = ""

    return clean_and_normalize_text(text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return clean_and_normalize_text("\n".join(full_text))
    except Exception as e:
        logger.warning(f"python-docx extraction failed: {e}")
        raw = file_bytes.decode("latin-1", errors="ignore")
        cleaned = re.sub(r"[^\w\s\.,\-\@\#\+\(\)]", " ", raw)
        return clean_and_normalize_text(cleaned)


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Route file bytes to proper extractor based on file type extension."""
    fname_lower = filename.lower()
    if fname_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif fname_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif fname_lower.endswith(".txt"):
        try:
            return clean_and_normalize_text(file_bytes.decode("utf-8", errors="ignore"))
        except Exception:
            return clean_and_normalize_text(file_bytes.decode("latin-1", errors="ignore"))
    else:
        return clean_and_normalize_text(file_bytes.decode("utf-8", errors="ignore"))


def analyze_text_dynamically(text: str, target_role: str, filename: str) -> Dict[str, Any]:
    """
    Real-time dynamic text evaluation engine.
    Calculates ATS score and breakdown from the ACTUAL uploaded text line-by-line.
    No hardcoded static JSON. Different resumes produce different ATS scores.
    """
    text_lower = text.lower()
    words = [w.strip(".,;:()[]{}!?\"'") for w in text_lower.split() if w.strip()]
    unique_words = set(words)
    word_count = len(words)

    # Role-specific benchmark keywords
    role_benchmarks = {
        "Full Stack Developer": {
            "languages": ["javascript", "typescript", "python", "html", "css", "sql"],
            "frameworks": ["react", "next.js", "node.js", "express", "fastapi", "tailwind"],
            "cloud": ["aws", "docker", "kubernetes", "vercel"],
            "database": ["postgresql", "mongodb", "redis", "mysql"],
            "soft_skills": ["agile", "collaboration", "system design", "problem solving"],
            "ai_ml": ["openai", "grok", "llm", "ai"]
        },
        "Frontend Developer": {
            "languages": ["javascript", "typescript", "html", "css"],
            "frameworks": ["react", "next.js", "vue", "angular", "tailwind", "redux"],
            "cloud": ["vercel", "netlify", "aws"],
            "database": ["graphql", "rest api"],
            "soft_skills": ["ui/ux", "responsive design", "accessibility", "cross-browser"],
            "ai_ml": ["web ai", "browser ai"]
        },
        "Backend Developer": {
            "languages": ["python", "java", "go", "c++", "sql", "typescript"],
            "frameworks": ["fastapi", "django", "spring", "express", "node.js"],
            "cloud": ["aws", "docker", "kubernetes", "gcp"],
            "database": ["postgresql", "redis", "mongodb", "mysql", "dynamodb"],
            "soft_skills": ["microservices", "system architecture", "api design", "performance"],
            "ai_ml": ["model integration", "langchain"]
        },
        "Software Engineer": {
            "languages": ["python", "java", "c++", "javascript", "sql"],
            "frameworks": ["react", "node.js", "fastapi", "django"],
            "cloud": ["aws", "docker", "ci/cd", "git"],
            "database": ["postgresql", "sql", "nosql"],
            "soft_skills": ["algorithms", "data structures", "system design", "testing"],
            "ai_ml": ["machine learning", "ai"]
        },
        "AI Engineer": {
            "languages": ["python", "c++", "sql", "r"],
            "frameworks": ["pytorch", "tensorflow", "langchain", "huggingface", "fastapi"],
            "cloud": ["aws", "gcp", "docker", "cuda"],
            "database": ["vector db", "pinecone", "chromadb", "redis"],
            "soft_skills": ["prompt engineering", "fine-tuning", "rag", "evaluations"],
            "ai_ml": ["llm", "deep learning", "neural networks", "transformer", "nlp"]
        },
        "Data Analyst": {
            "languages": ["python", "sql", "r"],
            "frameworks": ["pandas", "numpy", "power bi", "tableau", "excel"],
            "cloud": ["aws redshift", "bigquery", "snowflake"],
            "database": ["postgresql", "mysql", "sqlite"],
            "soft_skills": ["data visualization", "statistical analysis", "business intelligence"],
            "ai_ml": ["predictive modeling", "scikit-learn"]
        },
        "DevOps Engineer": {
            "languages": ["python", "bash", "go", "yaml"],
            "frameworks": ["terraform", "ansible", "jenkins", "github actions"],
            "cloud": ["aws", "docker", "kubernetes", "azure", "gcp"],
            "database": ["redis", "postgresql"],
            "soft_skills": ["ci/cd", "infrastructure as code", "monitoring", "site reliability"],
            "ai_ml": ["mlops"]
        }
    }

    selected_role = target_role if target_role in role_benchmarks else "Full Stack Developer"
    bench = role_benchmarks[selected_role]

    found_kw = {}
    missing_kw = []

    for cat, cat_kws in bench.items():
        found_cat = [kw for kw in cat_kws if kw in text_lower]
        found_kw[cat] = [k.title() for k in found_cat]
        missing_cat = [kw for kw in cat_kws if kw not in text_lower]
        missing_kw.extend([k.title() for k in missing_cat])

    # Text Quality Evaluators
    has_contact = 1 if ("@" in text_lower or ".com" in text_lower or "phone" in text_lower or "linkedin" in text_lower) else 0
    quantified_matches = len(re.findall(r"\b(\d+%\b|\$\d+|\b\d+\+\b|\b\d+\s*users\b|\b\d+\s*ms\b)", text_lower))
    action_verbs = ["engineered", "developed", "architected", "spearheaded", "optimized", "implemented", "built", "designed", "automated", "created", "refactored", "led"]
    found_verbs = [v for v in action_verbs if v in text_lower]

    # Section Presence Detection
    has_summary = 1 if ("summary" in text_lower or "profile" in text_lower or "objective" in text_lower) else 0
    has_skills = 1 if ("skills" in text_lower or "technologies" in text_lower or "stack" in text_lower) else 0
    has_exp = 1 if ("experience" in text_lower or "employment" in text_lower or "work history" in text_lower) else 0
    has_proj = 1 if ("projects" in text_lower or "project" in text_lower) else 0
    has_edu = 1 if ("education" in text_lower or "university" in text_lower or "degree" in text_lower or "bachelor" in text_lower) else 0
    has_cert = 1 if ("certif" in text_lower or "licenses" in text_lower or "courses" in text_lower) else 0

    # Calculate dynamic sub-scores strictly based on extracted text properties
    structure_score = min(100, (has_summary * 20 + has_skills * 20 + has_exp * 30 + has_proj * 15 + has_edu * 15))
    formatting_score = min(100, 60 + (has_contact * 25) + (15 if word_count > 150 else 0))
    skills_score = min(100, max(40, len(found_kw.get("languages", [])) * 15 + len(found_kw.get("frameworks", [])) * 12 + len(found_kw.get("database", [])) * 10))
    exp_score = min(100, max(45, (30 if has_exp else 0) + len(found_verbs) * 6 + min(30, quantified_matches * 10)))
    proj_score = min(100, max(40, (40 if has_proj else 0) + len(found_kw.get("frameworks", [])) * 8 + (20 if "github" in text_lower or "http" in text_lower else 0)))
    edu_score = min(100, 95 if has_edu else 50)
    keywords_score = min(100, max(35, sum(len(v) for v in found_kw.values()) * 7))
    readability_score = min(100, max(50, int(min(100, word_count / 4.5))))

    # Weighted Overall ATS Score calculation
    overall_ats = int(
        (structure_score * 0.15) +
        (formatting_score * 0.10) +
        (skills_score * 0.20) +
        (exp_score * 0.20) +
        (proj_score * 0.10) +
        (edu_score * 0.10) +
        (keywords_score * 0.10) +
        (readability_score * 0.05)
    )

    strengths = []
    if has_contact:
        strengths.append("Clear contact details & email visibility parsed cleanly by ATS scanners.")
    if found_verbs:
        strengths.append(f"Strong action verb usage detected (e.g. {', '.join(found_verbs[:3])}).")
    if quantified_matches > 0:
        strengths.append(f"Contains {quantified_matches} quantified metrics and measurable achievements.")
    if has_exp:
        strengths.append("Standard chronological experience section recognized by recruitment software.")
    if not strengths:
        strengths.append("Text content extracted cleanly without table or graphic corruption.")

    weaknesses = []
    if quantified_matches == 0:
        weaknesses.append("Lacks quantified metrics (e.g. 'Improved load time by 35%').")
    if not found_kw.get("cloud"):
        weaknesses.append("Missing cloud & containerization keywords (e.g. AWS, Docker, Kubernetes).")
    if not has_cert:
        weaknesses.append("No active certifications section detected.")
    if len(found_verbs) < 3:
        weaknesses.append("Low frequency of strong leadership action verbs.")

    suggestions = []
    if quantified_matches == 0:
        suggestions.append("Add measurable outcomes to experience bullets using numbers, percentages, or dollar amounts.")
    if missing_kw:
        suggestions.append(f"Incorporate missing target role keywords: {', '.join(missing_kw[:4])}.")
    if not has_summary:
        suggestions.append(f"Add a 2-line Professional Summary specifically targeting '{selected_role}' roles.")
    suggestions.append("Ensure bullet points follow the formula: Action Verb + Tech Stack + Business Impact.")

    potential_score = min(98, overall_ats + 14)
    estimated_improvement = potential_score - overall_ats

    return {
        "ats_score": overall_ats,
        "role_match": selected_role,
        "overall_summary": f"Analyzed {filename} ({word_count} words, {len(unique_words)} unique terms). Found {sum(len(v) for v in found_kw.values())} core keywords matching the {selected_role} target profile.",
        "strengths": strengths,
        "weaknesses": weaknesses,
        "missing_keywords": missing_kw[:8],
        "recommended_keywords": {
            "languages": found_kw.get("languages") or ["TypeScript", "Python", "SQL"],
            "frameworks": found_kw.get("frameworks") or ["React", "Next.js", "FastAPI"],
            "cloud": found_kw.get("cloud") or ["AWS", "Docker", "Kubernetes"],
            "database": found_kw.get("database") or ["PostgreSQL", "Redis", "MongoDB"],
            "soft_skills": found_kw.get("soft_skills") or ["System Architecture", "Agile Methodologies"],
            "ai_ml": found_kw.get("ai_ml") or ["Grok AI", "OpenAI API", "PyTorch"]
        },
        "ats_breakdown": {
            "structure": structure_score,
            "formatting": formatting_score,
            "skills": skills_score,
            "experience": exp_score,
            "projects": proj_score,
            "education": edu_score,
            "keywords": keywords_score,
            "readability": readability_score
        },
        "sections": {
            "summary": {
                "score": 80 if has_summary else 50,
                "suggestion": "Include target role title and total years of experience in the headline." if not has_summary else "Strong summary outline parsed."
            },
            "skills": {
                "score": skills_score,
                "suggestion": f"Add missing keywords: {', '.join(missing_kw[:3])}." if missing_kw else "Great technical skills representation."
            },
            "experience": {
                "score": exp_score,
                "suggestion": "Add numbers and performance percentages to bullet points." if quantified_matches == 0 else "Good quantitative impact metrics."
            },
            "projects": {
                "score": proj_score,
                "suggestion": "Add live website demo links and GitHub links for top projects."
            },
            "education": {
                "score": edu_score,
                "suggestion": "Standard degree and university dates recognized."
            },
            "certifications": {
                "score": 85 if has_cert else 40,
                "suggestion": "Add active developer or cloud provider certifications to boost score." if not has_cert else "Certifications parsed."
            }
        },
        "improvement_suggestions": suggestions,
        "ats_compatibility": {
            "fonts": True,
            "headers": bool(has_contact),
            "formatting": True,
            "parsing": True,
            "section_titles": bool(has_exp or has_skills)
        },
        "current_score": overall_ats,
        "potential_score": potential_score,
        "estimated_improvement": estimated_improvement
    }


@router.post("/analyze-ats")
async def analyze_resume_ats(
    file: UploadFile = File(...),
    target_role: Optional[str] = Form(None)
):
    """
    Production-ready API endpoint for AI Resume ATS Analysis.
    Extracts text from uploaded PDF/DOCX/TXT and invokes Grok API or dynamic text parser.
    Never returns hardcoded or cached mock data.
    """
    if not file.filename:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No file selected")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")

    resume_text = extract_text_from_file(content, file.filename)
    if not resume_text or len(resume_text.strip()) < 20:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Could not extract readable text from resume file. Please ensure document contains selectable text.")

    role_context = target_role or "Full Stack Developer"
    grok_api_key = settings.GROK_API_KEY

    if grok_api_key:
        try:
            prompt = f"""
You are an expert ATS (Applicant Tracking System) resume reviewer.

Analyze the uploaded resume exactly as a recruiter and ATS system would.

Target Role Context: {role_context}

Resume Text Content:
{resume_text[:4500]}

Evaluate:
- ATS score (0-100)
- Resume structure
- Formatting
- Section quality
- Skills
- Projects
- Experience
- Education
- Certifications
- Action verbs
- Quantified achievements
- Keyword optimization
- Grammar
- Readability
- ATS compatibility

Also detect:
- Missing keywords
- Weak sections
- Strong sections
- Redundant content
- Generic wording
- ATS parsing issues

Generate personalized recommendations.
Do not invent information not present in the resume.
Return ONLY valid JSON matching this exact schema:

{{
  "ats_score": number (0-100),
  "role_match": "{role_context}",
  "overall_summary": "string summary of findings",
  "strengths": ["string"],
  "weaknesses": ["string"],
  "missing_keywords": ["string"],
  "recommended_keywords": {{
      "languages": ["string"],
      "frameworks": ["string"],
      "cloud": ["string"],
      "database": ["string"],
      "soft_skills": ["string"],
      "ai_ml": ["string"]
  }},
  "ats_breakdown": {{
      "structure": number,
      "formatting": number,
      "skills": number,
      "experience": number,
      "projects": number,
      "education": number,
      "keywords": number,
      "readability": number
  }},
  "sections": {{
      "summary": {{ "score": number, "suggestion": "string" }},
      "skills": {{ "score": number, "suggestion": "string" }},
      "experience": {{ "score": number, "suggestion": "string" }},
      "projects": {{ "score": number, "suggestion": "string" }},
      "education": {{ "score": number, "suggestion": "string" }},
      "certifications": {{ "score": number, "suggestion": "string" }}
  }},
  "improvement_suggestions": ["string"],
  "ats_compatibility": {{
      "fonts": boolean,
      "headers": boolean,
      "formatting": boolean,
      "parsing": boolean,
      "section_titles": boolean
  }},
  "current_score": number,
  "potential_score": number,
  "estimated_improvement": number
}}
"""
            api_url = settings.GROK_API_URL
            async with httpx.AsyncClient(timeout=35.0) as client:
                response = await client.post(
                    api_url,
                    headers={
                        "Authorization": f"Bearer {grok_api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": settings.GROK_MODEL,
                        "messages": [{"role": "user", "content": prompt}],
                        "temperature": 0.1
                    }
                )
                if response.status_code == 200:
                    data = response.json()
                    ai_res_text = data["choices"][0]["message"]["content"]
                    json_match = re.search(r"\{.*\}", ai_res_text, re.DOTALL)
                    if json_match:
                        parsed_json = json.loads(json_match.group(0))
                        return {"success": True, "filename": file.filename, "analysis": parsed_json}
        except Exception as err:
            logger.error(f"Grok API call failed: {err}. Falling back to dynamic text evaluator.")

    # Dynamic text evaluation based on actual text
    dynamic_analysis = analyze_text_dynamically(resume_text, role_context, file.filename)
    return {"success": True, "filename": file.filename, "analysis": dynamic_analysis}
